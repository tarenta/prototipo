from docx import *
import os
from datetime import datetime
import locale

def geradoc():
    modelo_documento = 'declaração de hipossuficiência' #Aqui a variável deve receber o input da escolha do usuário sobre qual documento pretende gerar (essa linha pode nem ser necessária, podendo a variável correspondente ao inpu ficar diretamente na linha abaixo)

    documento = Document(f"Modelos de documentos\{modelo_documento}.docx")

    locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')
    hoje = datetime.now().strftime("%d de %B de %Y")

    nome_cliente1 = 'Nome do cliente'
    sexo_cliente1 = "Feminino"
    nacionalidade_cliente1 = "brasileira"
    estadocivil_cliente1 = "solteira"
    profissao_cliente1 = "professora"
    cpf_cliente1 = str("058.252.897-63")
    endereco_cliente1 = "Rua Dona Romana, nº 406, apto. 102, Engenho Novo, Rio de Janeiro - RJ, CEP 20710-200"
    nome_advogado1 = "Wallace Tarenta de Mendonça"
    sexo_advogado1 = "Masculino"
    nacionalidade_advogado1 = "Brasileiro"
    estadocivil_advogado1 = "solteiro"
    oab_advogado1 = str("248.898")
    email_advogado1 = "wallace@tarenta.com.br"
    endereco_advogado1 = "Rua Santo Irineu, 211, Campo Grande, Rio de Janeiro - RJ, CEP 23082-470"
    nome_terceiro1 = "Fulano de tal"
    sexo_terceiro1 = "Masculino"
    nacionalidade_terceiro1 = "Brasileira"
    estadocivil_terceiro1 = "Casado"
    profissao_terceiro1 = "Soldador"
    cpf_terceiro1 = "175.985.632-22"
    endereco_terceiro1 = "Rua do Igarapé, 1.060, Vale do Ambi, Duque de Caxias - RJ, CEP 56847-885"
    datahoje = hoje.lower()

    #Eu sei que tem um jeito melhor de fazer isso (até porque preciso lidar com os artigos no plural e outras flexões, como do, da, dos, das etc.)
    if sexo_cliente1 == "Masculino":
        artigo_cliente1 = "o"
    else: artigo_cliente1 = "a"

    if sexo_advogado1 == "Masculino":
        artigo_advogado1 = "o"
    else: artigo_advogado1 = "a"

    if sexo_terceiro1 == "Masculino":
        artigo_terceiro1 = "o"
    else: artigo_terceiro1 = "a"

    referencias = {
        "#nome_cliente1": nome_cliente1,
        "#sexo_cliente1": sexo_cliente1,
        "#nacionalidade_cliente1": nacionalidade_cliente1,
        "#estadocivil_cliente1": estadocivil_cliente1,
        "#profissão_cliente1": profissao_cliente1,
        "#cpf_cliente1": cpf_cliente1,
        "#endereco_cliente1": endereco_cliente1,
        "#artigo_cliente1": artigo_cliente1,
        "#nome_advogado1": nome_advogado1,
        "#sexo_advogado1": sexo_advogado1,
        "#nacionalidade_advogado1": nacionalidade_advogado1,
        "#estadocivil_advogado1": estadocivil_advogado1,
        "#oab_advogado1": oab_advogado1,
        "#email_advogado1": email_advogado1,
        "#endereco_advogado1": endereco_advogado1,
        "#artigo_advogado1": artigo_advogado1,
        "#nome_terceiro1": nome_terceiro1,
        "#genero_terceiro1": sexo_terceiro1,
        "#nacionalidade_terceiro1": nacionalidade_terceiro1,
        "#estadocivil_terceiro1": estadocivil_terceiro1,
        "#profissão_terceiro1": profissao_terceiro1,
        "#cpf_terceiro1": cpf_terceiro1,
        "#endereco_terceiro1": endereco_terceiro1,
        "#artigo_terceiro1": artigo_terceiro1,
        "#datahoje": datahoje,
    }

    for paragrafo in documento.paragraphs:
        runs_copy = paragrafo.runs.copy()
        for run in runs_copy:
            for chave in referencias:
                if chave in run.text:
                    substituto = referencias[chave]
                    run.text = run.text.replace(chave, str(substituto))

    for table in documento.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragrafo in cell.paragraphs:
                    runs_copy = paragrafo.runs.copy()
                    for run in runs_copy:
                        for chave in referencias:
                            if chave in run.text:
                                substituto = referencias[chave]
                                run.text = run.text.replace(chave, str(substituto))


    documento.save("resultado.docx") 
    os.startfile("resultado.docx")
