from tkinter import *
from tkinter import ttk
from docx import *
import os
from datetime import datetime
import locale
    
tomfundo1 = '#161B26'  # azul escuro
tomfundo2 = '#3F3831'  # Marrom escuro
tomfundo3 = '#897B6B'  # Marrom claro
tomfundo4 = '#BCAE98'  # Bege escuro
tomfundo5 = '#D7CBB9'  # Bege claro
tomclaro1 = '#686868'  # Cinza escuro
tomclaro2 = '#BFBFC1'  # Cinza claro
tomclaro3 = '#F5F5F5'  # Branco
destaque1 = '#317888'  # Azul
destaque2 = '#C78C00'  # Dourado

modelo_documento = ''

def geradoc():
    documento = Document(f"Modelos de documentos\{modelo_documento}.docx")

    locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')
    hoje = datetime.now().strftime("%d de %B de %Y")

    datahoje = hoje.lower()
    sexo_advogado1 = 'Masculino'

    if cbox_sexo_cliente1.get() == "Masculino":
        artigo_cliente1 = "o"
    else: artigo_cliente1 = "a"

    if sexo_advogado1 == "Masculino":
        artigo_advogado1 = "o"
    else: artigo_advogado1 = "a"

    if cbox_sexo_terceiro1.get() == "Masculino":
        artigo_terceiro1 = "o"
    else: artigo_terceiro1 = "a"

    referencias = {
        "#nome_cliente1": entry_nome_cliente1.get(),
        "#sexo_cliente1": cbox_sexo_cliente1.get(),
        "#nacionalidade_cliente1": cbox_nacionalidade_cliente1.get(),
        "#estadocivil_cliente1": cbox_estadocivil_cliente1.get(),
        "#profissão_cliente1": entry_profissao_cliente1.get(),
        "#cpf_cliente1": str(entry_cpf_cliente1.get()),
        "#endereco_cliente1": f"{cbox_tipoEndereço_cliente1.get()} {entry_logradouroEndereco_cliente1.get()}, {entry_numeroEndereco_cliente1.get()}, {entry_complementoEndereco_cliente1.get()}, {entry_bairroEndereco_cliente1.get()}, {entry_cidadeEndereco_cliente1.get()} - {cbox_estadoEndereço_cliente1.get()}, CEP: {entry_cepEndereco_cliente1.get()}",
        "#artigo_cliente1": artigo_cliente1,
        "#nome_advogado1": 'Wallace Tarenta de Mendonça',
        "#sexo_advogado1": sexo_advogado1,
        "#artigo_advogado1": artigo_advogado1,
        "#nacionalidade_advogado1": 'brasileiro',
        "#estadocivil_advogado1": 'solteiro',
        "#oab_advogado1": str('145.453.407-92'),
        "#email_advogado1": 'wallace@tarenta.com.br',
        "#endereco_advogado1": 'Campo Grande, Rio de Janeiro',
        "#nome_terceiro1": entry_nome_terceiro1.get(),
        "#sexo_terceiro1": cbox_sexo_terceiro1.get(),
        "#artigo_terceiro1": artigo_terceiro1,
        "#nacionalidade_terceiro1": cbox_nacionalidade_terceiro1.get(),
        "#estadocivil_terceiro1": cbox_estadocivil_terceiro1.get(),
        "#profissão_terceiro1": entry_profissao_terceiro1.get(),
        "#cpf_terceiro1": str(entry_cpf_terceiro1.get()),
        "#endereco_terceiro1": "Rua do Igarapé, 1.060, Vale do Ambi, Duque de Caxias - RJ, CEP 56847-885",
        "#artigo_terceiro1": artigo_terceiro1,
        "#datahoje": datahoje
    }

    for paragrafo in documento.paragraphs:
        runs_copy = paragrafo.runs.copy()
        for run in runs_copy:
            for chave in referencias:
                if chave in run.text:
                    substituto = referencias[chave]
                    run.text = run.text.replace(chave, str(substituto))

    for table in documento.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragrafo in cell.paragraphs:
                    runs_copy = paragrafo.runs.copy()
                    for run in runs_copy:
                        for chave in referencias:
                            if chave in run.text:
                                substituto = referencias[chave]
                                run.text = run.text.replace(chave, str(substituto))


    documento.save("resultado.docx") 
    os.startfile("resultado.docx")

def seleciona_documento(valor):
   global modelo_documento
   modelo_documento = valor

def show_frame(frame):
    frame.tkraise()

def titulo_padrao(parent, texto):
    frame = Frame(parent, bg=tomfundo1)
    frame.rowconfigure(0, weight=1)
    frame.columnconfigure(0, weight=1)
    frame.grid(row=0, column=0, columnspan=2, sticky='nsew')
    texto = Label(frame, text=f'{texto}', bg='white')
    texto.grid(row=0, column=0)
    return frame

def menu_padrao(parent):
    frame = Frame(parent, bg=tomfundo1)
    frame.columnconfigure(0, weight=1, uniform='a')
    frame.rowconfigure((0, 1, 2, 3, 4, 5, 6, 7, 8, 9), weight=1, uniform='a')
    frame.grid(row=1, column=0, sticky='nsew')
    return frame

def painel_padrao(parent):
    frame = Frame(parent, bg=tomclaro1)
    frame.columnconfigure((0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11,
                          12, 13, 14, 15, 16, 17, 18, 19), weight=1, uniform='a')
    frame.rowconfigure((0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11,
                       12, 13, 14, 15, 16, 17, 18, 19), weight=1, uniform='a')
    frame.grid(row=1, column=1, sticky='nsew')
    return frame

def botoes_menu(parent):
    botao_inicio = Button(
        parent, background=destaque1, foreground=tomfundo1, activebackground=destaque2, activeforeground=tomfundo2, highlightthickness=2, highlightbackground='orange', highlightcolor='white', font=('Arial', 16, 'bold'), text='Início', command=lambda: show_frame(janela_home))
    botao_inicio.grid(row=2, column=0, padx=(0, 10), pady=5, sticky='nsew')

    botao_gerador = Button(
        parent, background=destaque1, foreground=tomfundo1, activebackground=destaque2, activeforeground=tomfundo2, highlightthickness=2, highlightbackground='orange', highlightcolor='white', font=('Arial', 16, 'bold'), text='Gerador', command=lambda: show_frame(janela_gerador))
    botao_gerador.grid(row=3, column=0, padx=(0, 10), pady=5, sticky='nsew')

    botao_acompanhamentoprocessual = Button(
        parent, background=destaque1, foreground=tomfundo1, activebackground=destaque2, activeforeground=tomfundo2, highlightthickness=2, highlightbackground='orange', highlightcolor='white', font=('Arial', 16, 'bold'), text='Acompanhamento Processual', command=lambda: show_frame(janela_acompanhamentoprocessual))
    botao_acompanhamentoprocessual.grid(row=4, column=0, padx=(0, 10), pady=5, sticky='nsew')

    botao_diariooficial = Button(
        parent, background=destaque1, foreground=tomfundo1, activebackground=destaque2, activeforeground=tomfundo2, highlightthickness=2, highlightbackground='orange', highlightcolor='white', font=('Arial', 16, 'bold'), text='Diário Oficial', command=lambda: show_frame(janela_diariooficial))
    botao_diariooficial.grid(row=5, column=0, padx=(0, 10), pady=5, sticky='nsew')

    botao_jurimetria = Button(
        parent, background=destaque1, foreground=tomfundo1, activebackground=destaque2, activeforeground=tomfundo2, highlightthickness=2, highlightbackground='orange', highlightcolor='white', font=('Arial', 16, 'bold'), text='Jurimetria', command=lambda: show_frame(janela_jurimetria))
    botao_jurimetria.grid(row=6, column=0, padx=(0, 10), pady=5, sticky='nsew')

    botao_jurisprudencia = Button(
        parent, background=destaque1, foreground=tomfundo1, activebackground=destaque2, activeforeground=tomfundo2, highlightthickness=2, highlightbackground='orange', highlightcolor='white', font=('Arial', 16, 'bold'), text='Jurisprudência', command=lambda: show_frame(janela_jurisprudencia))
    botao_jurisprudencia.grid(row=7, column=0, padx=(0, 10), pady=5, sticky='nsew')

    return botao_inicio, botao_gerador, botao_acompanhamentoprocessual, botao_diariooficial, botao_jurimetria, botao_jurisprudencia

def botao_proximo(parent, destino):
    nome = Button(parent, text='Próximo', command=lambda: show_frame(destino))
    nome.grid(row=18, column=17, sticky='new')
    return nome

def botao_voltar(parent, destino):
    nome = Button(parent, text='Voltar', command=lambda: show_frame(destino))
    nome.grid(row=18, column=2, sticky='new')

def labelentry(parent, labelTexto, nRow, nColumn, entryRowSpan, entryColumnSpan, npadx, npady):
    labelNome = Label(parent, text=str(labelTexto), bg=parent.cget("bg"))
    labelNome.grid(row=int(nRow), column=int(nColumn),
                   padx=npadx, pady=npady, sticky='sw')
    entryNome = Entry(parent)
    entryNome.grid(row=int(nRow + 1), column=int(nColumn),
                   rowspan=int(entryRowSpan), columnspan=int(entryColumnSpan), padx=npadx, pady=npady, sticky='nsew')
    return entryNome

def labelcombobox(parent, labelTexto, cboxValues, nRow, nColumn, cboxRowSpan, cboxColumnSpan, npadx, npady):
    labelNome = Label(parent, text=str(labelTexto), bg=parent.cget("bg"))
    labelNome.grid(row=int(nRow), column=int(nColumn),
                   padx=npadx, pady=npady, sticky='sw')
    cboxNome = ttk.Combobox(parent, values=cboxValues)
    cboxNome.grid(row=int(nRow + 1), column=int(nColumn),
                  rowspan=int(cboxRowSpan), columnspan=int(cboxColumnSpan), padx=npadx, pady=npady, sticky='nsew')
    return cboxNome

janela = Tk()

janela.state('zoomed')
janela.rowconfigure(0, weight=1)
janela.columnconfigure(0, weight=1)

# ====================== Criação das janelas padrão
janela_home = Frame(janela)
janela_gerador = Frame(janela)
janela_gerador_documentos = Frame(janela)
janela_gerador_documentos_opcoes = Frame(janela)
janela_gerador_documentos_form = Frame(janela)
janela_acompanhamentoprocessual = Frame(janela)
janela_diariooficial = Frame(janela)
janela_jurimetria = Frame(janela)
janela_jurisprudencia = Frame(janela)


# ======================= Configuração das janelas padrão
for frame in janela.winfo_children():
    frame.columnconfigure(0, weight=1, uniform='a')
    frame.columnconfigure(1, weight=4, uniform='a')
    frame.rowconfigure(0, weight=1, uniform='a')
    frame.rowconfigure(1, weight=9, uniform='a')
    frame.grid(row=0, column=0, sticky='nsew')


# ======================= Configuração dos frames (titulo, menu e painel)

titulo_home = titulo_padrao(janela_home, 'início')
titulo_gerador = titulo_padrao(janela_gerador, 'Gerador')
titulo_gerador_documentos = titulo_padrao(janela_gerador_documentos, 'Escolha um documento')
titulo_gerador_documentos_opcoes = titulo_padrao(janela_gerador_documentos_opcoes, 'Opções do documento')
titulo_gerador_documentos_form = titulo_padrao(janela_gerador_documentos_form, 'Formulário')
titulo_acompanhamentoprocessual = titulo_padrao(janela_acompanhamentoprocessual, 'Acompanhamento Processual')
titulo_diariooficial = titulo_padrao(janela_diariooficial, 'Diário Oficial')
titulo_jurimetria = titulo_padrao(janela_jurimetria, 'Jurimetria')
titulo_jurisprudencia = titulo_padrao(janela_jurisprudencia, 'Jurisprudência')

menu_home = menu_padrao(janela_home)
menu_gerador = menu_padrao(janela_gerador)
menu_gerador_documentos = menu_padrao(janela_gerador_documentos)
menu_gerador_documentos_opcoes = menu_padrao(janela_gerador_documentos_opcoes)
menu_gerador_documentos_form = menu_padrao(janela_gerador_documentos_form)
menu_acompanhamentoprocessual = menu_padrao(janela_acompanhamentoprocessual)
menu_diariooficial = menu_padrao(janela_diariooficial)
menu_jurimetria = menu_padrao(janela_jurimetria)
menu_jurisprudencia = menu_padrao(janela_jurisprudencia)

painel_home = painel_padrao(janela_home)
painel_gerador = painel_padrao(janela_gerador)
painel_gerador_documentos = painel_padrao(janela_gerador_documentos)
painel_gerador_documentos_opcoes = painel_padrao(janela_gerador_documentos_opcoes)
painel_gerador_documentos_form = painel_padrao(janela_gerador_documentos_form)
painel_acompanhamentoprocessual = painel_padrao(janela_acompanhamentoprocessual)
painel_diariooficial = painel_padrao(janela_diariooficial)
painel_jurimetria = painel_padrao(janela_jurimetria)
painel_jurisprudencia = painel_padrao(janela_jurisprudencia)


# =========================== Configuração painéis
# for frame in (painel_home, painel_gerador, painel_gerador_documentos, painel_gerador_documentos_opcoes, painel_gerador_documentos_form):


# ========================== Criando botões das funções
botao_proximo_gerador = botao_proximo(painel_gerador, janela_gerador_documentos)
botao_proximo_gerador_documentos = botao_proximo(painel_gerador_documentos, janela_gerador_documentos_opcoes)
botao_proximo_gerador_documentos_opcoes = botao_proximo(painel_gerador_documentos_opcoes, janela_gerador_documentos_form)
botao_voltar_gerador_documentos = botao_voltar(painel_gerador_documentos, janela_gerador)
botao_voltar_gerador_documentos_opcoes = botao_voltar(painel_gerador_documentos_opcoes, janela_gerador_documentos)
botao_voltar_gerador_documentos_form = botao_voltar(painel_gerador_documentos_form, janela_gerador_documentos_opcoes)

# ========================== Home
titulo_home
menu_home
botoes_menu(menu_home)
painel_home

# ========================== Gerador
titulo_gerador
menu_gerador
botoes_menu(menu_gerador)

painel_gerador
quadro_gerador = LabelFrame(painel_gerador, text='O que você quer gerar?')
quadro_gerador.columnconfigure((0, 1), weight=1)
quadro_gerador.rowconfigure(0, weight=1)
quadro_gerador.grid(row=4, column=4, rowspan=12, columnspan=12, sticky='nsew')

botao_quadro_gerador_peticao = Button(quadro_gerador, text='Petição')
botao_quadro_gerador_peticao.grid(row=0, column=0)
botao_quadro_gerador_documento = Button(
    quadro_gerador, text='Documento', command=lambda: show_frame(janela_gerador_documentos))
botao_quadro_gerador_documento.grid(row=0, column=1)

# ========================== Gerador-documentos
titulo_gerador_documentos
menu_gerador_documentos
botoes_menu(menu_gerador_documentos)
painel_gerador_documentos
botao_voltar_gerador_documentos

quadro_gerador_documentos = LabelFrame(
    painel_gerador_documentos, text='O que você quer gerar?')
quadro_gerador_documentos.columnconfigure(0, weight=1)
quadro_gerador_documentos.rowconfigure((0, 1, 2), weight=1)
quadro_gerador_documentos.grid(
    row=4, column=4, rowspan=12, columnspan=12, sticky='nsew')

botao_quadro_gerador_documentos_procuracao = Button(
    quadro_gerador_documentos, text='Procuração', command=lambda: [show_frame(janela_gerador_documentos_form), seleciona_documento('procuração ad juditia')])
botao_quadro_gerador_documentos_procuracao.grid(row=0, column=0)
botao_quadro_gerador_documentos_declaraHipo = Button(
    quadro_gerador_documentos, text='Declaração de Hipossuficiência', command=lambda: [show_frame(janela_gerador_documentos_form), seleciona_documento('declaração de hipossuficiência')])
botao_quadro_gerador_documentos_declaraHipo.grid(row=1, column=0)
botao_quadro_gerador_documentos_declaraResidencia = Button(
    quadro_gerador_documentos, text='Declaração de Residência', command=lambda: [show_frame(janela_gerador_documentos_form), seleciona_documento('declaração de residência')])
botao_quadro_gerador_documentos_declaraResidencia.grid(row=2, column=0)

# ========================== Gerador-documentos-opções TODO: por enquanto eu pulei a criação dessa janela, mas ela será necessária para o usuário decidir questões como quantidade de autores/declarante/advogados e outras coisas
titulo_gerador_documentos_opcoes
menu_gerador_documentos_opcoes
botoes_menu(menu_gerador_documentos_opcoes)
painel_gerador_documentos_opcoes
botao_voltar_gerador_documentos_opcoes

# ========================== Gerador-documentos-formulário
titulo_gerador_documentos_form
menu_gerador_documentos_form
botoes_menu(menu_gerador_documentos_form)
painel_gerador_documentos_form
botao_voltar_gerador_documentos_form

quadro_dados_cliente1 = LabelFrame(
    painel_gerador_documentos_form, text='Dados do Declarante', bg=tomclaro2)
quadro_dados_cliente1.columnconfigure(
    (1, 2, 3, 4, 5), weight=3, uniform='a')
quadro_dados_cliente1.columnconfigure(
    (0, 6), weight=1, uniform='a')
quadro_dados_cliente1.rowconfigure(
    (0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15), weight=1, uniform='a')
quadro_dados_cliente1.grid(
    row=1, column=6, rowspan=8, columnspan=8, sticky='nsew')

entry_nome_cliente1 = labelentry(quadro_dados_cliente1, 'Nome completo', 0, 1, 1, 3, (0, 10), 0)
entry_cpf_cliente1 = labelentry(quadro_dados_cliente1, 'CPF', 0, 4, 1, 1, (10, 10), 0)
entry_rg_cliente1 = labelentry(quadro_dados_cliente1, 'RG', 0, 5, 1, 1, (10, 0), 0)
entry_email_cliente1 = labelentry(quadro_dados_cliente1, 'E-mail', 3, 1, 1, 2, (0, 10), 0)
entry_profissao_cliente1 = labelentry(quadro_dados_cliente1, 'Profissão', 3, 3, 1, 2, (10, 10), 0)
cbox_estadocivil_cliente1 = labelcombobox(quadro_dados_cliente1, 'Estado Civil', ['', 'Solteiro', 'Casado', 'Divorciado', 'Em união estável'], 3, 5, 1, 1, (10, 0), 0)
cbox_nacionalidade_cliente1 = labelcombobox(quadro_dados_cliente1, 'Nacionalidade', ['', 'Brasileiro', 'Estrangeiro'], 6, 1, 1, 1, (0, 10), 0)
cbox_sexo_cliente1 = labelcombobox(quadro_dados_cliente1, 'Sexo', ['', 'Masculino', 'Feminino', 'Outro', 'prefiro não responder'], 6, 2, 1, 1, (10, 10), 0)

quadro_endereco_cliente1 = LabelFrame(
    quadro_dados_cliente1, text='Endereço do declarante', bg=tomclaro2)
quadro_endereco_cliente1.columnconfigure(
    (0, 1, 2, 3, 4, 5), weight=1, uniform='a')
quadro_endereco_cliente1.rowconfigure((0, 1, 2, 3, 4), weight=2, uniform='a')
quadro_endereco_cliente1.rowconfigure(5, weight=1, uniform='a')
quadro_endereco_cliente1.grid(
    row=9, column=1, rowspan=6, columnspan=5, sticky='nsew')

cbox_tipoEndereço_cliente1 = labelcombobox(quadro_endereco_cliente1, 'Tipo', ['', 'Rua', 'Avenida', 'Estrada', 'Alameda', 'Avenida', 'Praça', 'Largo', 'Travessa'], 0, 0, 1, 1, (10, 10), 0)
entry_logradouroEndereco_cliente1 = labelentry(quadro_endereco_cliente1, 'Logradouro', 0, 1, 1, 4, (10, 10), 0)
entry_numeroEndereco_cliente1 = labelentry(quadro_endereco_cliente1, 'Número', 0, 5, 1, 1, (10, 10), 0)
entry_complementoEndereco_cliente1 = labelentry(quadro_endereco_cliente1, 'Complemento', 3, 0, 1, 1, (10, 10), 0)
entry_bairroEndereco_cliente1 = labelentry(quadro_endereco_cliente1, 'Bairro', 3, 1, 1, 2, (10, 10), 0)
entry_cidadeEndereco_cliente1 = labelentry(quadro_endereco_cliente1, 'Cidade', 3, 3, 1, 1, (10, 10), 0)
cbox_estadoEndereço_cliente1 = labelcombobox(quadro_endereco_cliente1, 'Estado', [" ", "AC", "AL", "AP", "AM", "BA", "CE", "DF", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR", "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO"],  3, 4, 1, 1, (10, 10), 0)
entry_cepEndereco_cliente1 = labelentry(quadro_endereco_cliente1, 'CEP', 3, 5, 1, 1, (10, 10), 0)

quadro_dados_terceiro1 = LabelFrame(
    painel_gerador_documentos_form, text='Dados do Declarado', bg= tomclaro2)
quadro_dados_terceiro1.columnconfigure(
    (1, 2, 3, 4, 5), weight=3, uniform='a')
quadro_dados_terceiro1.columnconfigure(
    (0, 6), weight=1, uniform='a')
quadro_dados_terceiro1.rowconfigure(
    (0, 1, 2, 3, 4, 5, 6, 7, 8), weight=1, uniform='a')
quadro_dados_terceiro1.grid(
    row=10, column=6, rowspan=6, columnspan=8, sticky='nsew')

entry_nome_terceiro1 = labelentry(quadro_dados_terceiro1, 'Nome completo', 0, 1, 1, 3, (0, 10), 0)
entry_cpf_terceiro1 = labelentry(quadro_dados_terceiro1, 'CPF', 0, 4, 1, 1, (10, 10), 0)
entry_rg_terceiro1 = labelentry(quadro_dados_terceiro1, 'RG', 0, 5, 1, 1, (10, 0), 0)
entry_email_terceiro1 = labelentry(quadro_dados_terceiro1, 'E-mail', 3, 1, 1, 2, (0, 10), 0)
entry_profissao_terceiro1 = labelentry(quadro_dados_terceiro1, 'Profissão', 3, 3, 1, 2, (10, 10), 0)
cbox_estadocivil_terceiro1 = labelcombobox(quadro_dados_terceiro1, 'Estado Civil', ['', 'Solteiro', 'Casado', 'Divorciado', 'Em união estável'], 3, 5, 1, 1, (10, 0), 0)
cbox_nacionalidade_terceiro1 = labelcombobox(quadro_dados_terceiro1, 'Nacionalidade', ['', 'Brasileiro', 'Estrangeiro'], 6, 1, 1, 1, (0, 10), 0)
cbox_sexo_terceiro1 = labelcombobox(quadro_dados_terceiro1, 'Sexo', ['', 'Masculino', 'Feminino', 'Outro', 'prefiro não responder'], 6, 2, 1, 1, (10, 10), 0)

botao_gerar = Button(painel_gerador_documentos_form, text='Criar Documento', command=geradoc)
botao_gerar.grid(row=17, column=9, columnspan=2, sticky='nsew')

# ========================== Acompanhamento Processual
titulo_acompanhamentoprocessual
menu_acompanhamentoprocessual
botoes_menu(menu_acompanhamentoprocessual)
painel_acompanhamentoprocessual

# ========================== Diário Oficial
titulo_diariooficial
menu_diariooficial
botoes_menu(menu_diariooficial)
painel_diariooficial

# ========================== Jurimetria
titulo_jurimetria
menu_jurimetria
botoes_menu(menu_jurimetria)
painel_jurimetria

# ========================== Jurisprudência
titulo_jurisprudencia
menu_jurisprudencia
botoes_menu(menu_jurisprudencia)
painel_jurisprudencia

show_frame(janela_home)

janela.mainloop()